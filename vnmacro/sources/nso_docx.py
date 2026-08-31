"""Extract indicators from the NSO narrative report (.docx).

Some of the most policy-relevant numbers never reach the statistical tables —
state budget revenue and expenditure, its composition, execution against the
annual plan, disbursed FDI, and (in quarterly issues) GDP growth and its
demand-side split all live only in prose:

    "Tổng thu ngân sách Nhà nước tháng Bảy ước đạt 259,7 nghìn tỷ đồng.
     Lũy kế ... bảy tháng năm 2026 ước đạt 1.834,6 nghìn tỷ đồng, bằng 72,5%
     dự toán năm và tăng 16,0% so với cùng kỳ năm trước."

So the patterns live in ``config/narrative_patterns.yaml`` as regexes with
named groups, keeping the brittle part out of the code where it can be fixed
without touching the pipeline.
"""
from __future__ import annotations

import datetime as dt
import logging
import re
from pathlib import Path

import docx

from ..config import load_yaml
from ..util import norm_ws, to_float

log = logging.getLogger(__name__)

_PATTERNS_CACHE: list[dict] | None = None


#: OLE2 compound-document magic — a genuine Word 97-2003 .doc.
OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

#: Control characters that pepper the raw Word stream between text runs.
_CTRL = r"\x00-\x08\x0b\x0c\x0e-\x1f�"
_RUN = re.compile(rf"[^{_CTRL}]{{40,}}")
_VIETNAMESE = re.compile(r"[àáâãèéêìíòóôõùúýăđĩũơưạ-ỹ]", re.I)


def _read_legacy_doc(path: Path) -> list[str]:
    """Recover text from a Word 97 .doc, which python-docx cannot open.

    Reports up to about 2017 ship binary .doc. Rather than take a dependency
    on LibreOffice, the WordDocument stream is decoded as UTF-16LE (which is
    how Word stores Vietnamese) and the long, diacritic-bearing runs are kept.

    This is deliberately crude, and safe *because* it is: the narrative
    patterns are long and specific ("Tổng thu ngân sách Nhà nước tháng ... ước
    đạt N nghìn tỷ đồng"), so damaged text fails to match rather than matching
    something wrong.
    """
    import olefile

    with olefile.OleFileIO(str(path)) as ole:
        if not ole.exists("WordDocument"):
            return []
        data = ole.openstream("WordDocument").read()

    text = data.decode("utf-16-le", errors="replace")
    runs = [norm_ws(r) for r in _RUN.findall(text) if _VIETNAMESE.search(r)]
    return [r for r in runs if r]


def read_text(path: Path) -> list[str]:
    """All paragraphs plus table cells, in document order, whitespace-normalised."""
    with path.open("rb") as fh:
        magic = fh.read(8)
    if magic == OLE2_MAGIC:
        return _read_legacy_doc(path)

    d = docx.Document(str(path))
    out = [norm_ws(p.text) for p in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [norm_ws(c.text) for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return [t for t in out if t]


def load_patterns() -> list[dict]:
    global _PATTERNS_CACHE
    if _PATTERNS_CACHE is None:
        cfg = load_yaml("narrative_patterns.yaml")
        pats = cfg.get("patterns", [])
        for p in pats:
            p["_re"] = re.compile(p["pattern"], re.I)
        _PATTERNS_CACHE = pats
    return _PATTERNS_CACHE


#: Words that flip the sign of the number that follows them.
NEGATIVE_WORDS = re.compile(r"gi[ảa]m|nh[ậa]p si[êe]u|th[âa]m h[ụu]t|b[ộo]i chi", re.I)


def _signed(value: float, direction: str | None) -> float:
    """'giảm 3,2%' -> -3.2; 'nhập siêu 20,52 tỷ USD' -> -20.52."""
    if direction and NEGATIVE_WORDS.search(direction.strip()):
        return -value
    return value


def extract(paragraphs: list[str], *, release: dict, raw_file: str) -> list[dict]:
    period = release["period_date"]
    month_date = release.get("month_date") or period
    if isinstance(month_date, str):
        month_date = dt.date.fromisoformat(month_date)
    if isinstance(period, str):
        period = dt.date.fromisoformat(period)
    vintage = release.get("wp_date", "")[:10]
    vintage = dt.date.fromisoformat(vintage) if vintage else None

    seen: set[str] = set()
    out: list[dict] = []
    for pat in load_patterns():
        for text in paragraphs:
            m = pat["_re"].search(text)
            if not m:
                continue
            gd = m.groupdict()
            val = to_float(gd.get("val"))
            if val is None:
                continue
            if pat.get("signed", False):
                val = _signed(val, gd.get("dir"))
            if pat.get("percent_change", False):
                # "tăng 16,0%" is published as a growth rate, keep it as one
                pass
            sid = pat["id"]
            if sid in seen:
                continue
            seen.add(sid)
            out.append({
                "series_id": f"NSO.{sid}",
                "dataset": "nso_narrative",
                "source": "NSO",
                "freq": release["freq"],
                # PHẢI là month_date, không phải period. Báo cáo quý ("Quý II và
                # sáu tháng đầu năm") mang số của THÁNG CUỐI quý, không phải
                # tháng đầu: giá trị luỹ kế trong đó là 6 tháng chứ không phải 4.
                # Gán vào tháng đầu quý thì chuỗi luỹ kế mất tính đơn điệu — số
                # "tháng 4" (thực ra 6 tháng) lớn hơn số tháng 5 thật, và mọi so
                # sánh theo mùa vụ đều sai kỳ. nso_xlsx.py đã dùng month_date;
                # chỗ này bị bỏ sót.
                "date": month_date,
                "ref_period": release["ref_period"],
                "value": val,
                "unit": pat.get("unit"),
                "scale": 0,
                "status": pat.get("status", "uoc_tinh"),
                "vintage": vintage,
                "partner": None,
                "breakdown": pat.get("breakdown"),
                "label_vi": pat.get("label_vi", sid),
                "measure": pat.get("scope", "period"),
                "dims": {"scope": pat.get("scope", "period"),
                         "matched": text[:400]},
                "raw_file": raw_file,
            })
            break        # first match in document order wins
    missing = [p["id"] for p in load_patterns() if p["id"] not in seen]
    if missing:
        log.info("%s: %d narrative patterns unmatched (%s)",
                 release.get("ref_period"), len(missing), ", ".join(missing[:6]))
    return out
